"""任务与编排器路由 — create/get/cancel 任务 + 后台编排器。"""

import asyncio
import logging
import re
import uuid
from pathlib import Path

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query
from fastapi.responses import FileResponse

from ..auth import GitHubUser, get_current_user
from ..config import get_settings
from ..services.session import get_session_manager
from .apikeys import _resolve_user_id, get_active_api_key
from .schemas.request import CreateTaskRequest
from .schemas.response import MessageResponse, TaskResponse

logger = logging.getLogger(__name__)

tasks_router = APIRouter()


# ── 附件内容提取（方案模式把上传的题目/数据文件读进问题上下文）────────

# 单个文件提取内容的上限（字符），避免超大文件撑爆上下文
_MAX_FILE_CHARS = 50000


def _extract_attachment_text(file_id: str, filename: str) -> str:
    """按扩展名提取已上传附件的文本内容。失败时返回提示信息而非抛错。"""
    settings = get_settings()
    uploads_dir = settings.project_root / "data" / "uploads"
    matches = list(uploads_dir.glob(f"{file_id}.*"))
    if not matches:
        return f"（附件 {filename} 未找到）"
    path = matches[0]
    suffix = path.suffix.lower()

    try:
        data = path.read_bytes()

        if suffix == ".pdf":
            # 解析函数定义在 knowledge_search_routes（god-files 拆分后迁移），
            # 从 knowledge_routes 导入会 ImportError 且被 except 静默吞掉
            from .knowledge_search_routes import _extract_pdf_text

            text = _extract_pdf_text(data)
        elif suffix == ".docx":
            from .knowledge_search_routes import _extract_docx_text

            text = _extract_docx_text(data)
        elif suffix in (".xlsx", ".xls", ".csv", ".tsv"):
            import io as _io

            import pandas as pd

            # 表格最多展示行数：超限截断并保留形状信息，避免大表撑爆上下文
            _MAX_DF_ROWS = 100

            def _df_to_text(df, title: str = "") -> str:
                """结构化输出：标题 + 形状（行×列）+ 前 N 行 markdown 表 + 截断提示。"""
                shape = f"{df.shape[0]} 行 × {df.shape[1]} 列"
                head = f"{title}【{shape}】" if title else f"【{shape}】"
                truncated = df.shape[0] > _MAX_DF_ROWS
                df_show = df.head(_MAX_DF_ROWS) if truncated else df
                try:
                    table = df_show.to_markdown(index=False)
                except Exception:
                    table = df_show.to_csv(index=False)
                tail = (
                    f"\n…（共 {df.shape[0]} 行，仅展示前 {_MAX_DF_ROWS} 行）"
                    if truncated
                    else ""
                )
                return f"{head}\n{table}{tail}"

            if suffix in (".csv", ".tsv"):
                sep = "\t" if suffix == ".tsv" else ","
                df = pd.read_csv(
                    _io.StringIO(data.decode("utf-8", errors="replace")),
                    sep=sep,
                )
                text = _df_to_text(df, title=filename)
            else:
                xls = pd.ExcelFile(_io.BytesIO(data))
                parts = []
                for sheet in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet)
                    parts.append(_df_to_text(df, title=f"工作表「{sheet}」"))
                text = "\n\n".join(parts)
        else:
            # txt / md / json / py / dat 等按文本读
            for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    text = data.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = data.decode("utf-8", errors="replace")

        text = (text or "").strip()
        if len(text) > _MAX_FILE_CHARS:
            text = text[:_MAX_FILE_CHARS] + f"\n…（内容过长，已截断至 {_MAX_FILE_CHARS} 字符）"
        return text or "（附件无可提取文本）"
    except HTTPException as e:
        return f"（附件 {filename} 解析失败: {e.detail}）"
    except Exception as e:  # noqa: BLE001
        logger.exception("附件提取失败 %s", filename)
        return f"（附件 {filename} 解析失败: {e}）"


def _build_problem_with_attachments(problem: str, files: list) -> str:
    """把附件内容拼接进问题描述，供分类/建模/求解各节点使用。"""
    if not files:
        return problem
    parts = [problem]
    for f in files:
        content = _extract_attachment_text(f.file_id, f.filename)
        parts.append(f"\n\n## 附件：{f.filename}\n{content}")
    return "".join(parts)


# ── Tasks ────────────────────────────────────────────────────────


@tasks_router.post("/tasks", response_model=TaskResponse)
async def create_task(
    req: CreateTaskRequest,
    background_tasks: BackgroundTasks,
    user: GitHubUser | None = Depends(get_current_user),
):
    """创建建模任务，后台启动编排器。"""
    # 检查是否有可用的 API Key（读 apikeys.json 文件 IO，不占主循环）
    uid = _resolve_user_id(user=user)
    active_key = await asyncio.to_thread(get_active_api_key, uid)
    if not active_key:
        raise HTTPException(
            status_code=400,
            detail="请先在首页配置你的 API Key，然后再提交任务。",
        )

    session_mgr = get_session_manager()
    # 把上传附件（题目PDF/数据Excel等）的内容提取进问题上下文，
    # 否则分类节点只能看到用户输入的一句话，无法理解题目。
    # PDF/Excel 解析是重 IO+CPU，同样丢线程池
    full_problem = await asyncio.to_thread(
        _build_problem_with_attachments, req.problem, req.files
    )
    task = session_mgr.create(problem=full_problem, mode=req.mode)
    task_id = task["task_id"]

    # 文件区：记录用户上传的附件（供前端文件区展示与下载）
    for f in req.files:
        session_mgr.add_artifact(
            task_id,
            {
                "type": "uploaded",
                "name": f.filename,
                "url": f"/api/files/{f.file_id}",
            },
        )

    # 初始化工作记忆（问题状态文档 + 检查点目录）
    # mkdir/write_text 属磁盘 IO，与上面同理不占主循环
    from app.services.working_memory import WorkingMemory

    await asyncio.to_thread(WorkingMemory(task_id).init_session, full_problem)

    # 在独立线程中运行编排器（节点含同步阻塞调用 llm.invoke / subprocess），
    # 以免阻塞事件循环导致 HTTP 响应体无法刷新、WS 进度卡住
    asyncio.create_task(
        asyncio.to_thread(_run_orchestrator_sync, task_id, full_problem, req.mode, uid)
    )

    return TaskResponse(**task)


@tasks_router.get("/tasks/{task_id}/status")
async def get_task_status(task_id: str, user: GitHubUser | None = Depends(get_current_user)):
    """查询方案模式任务的执行进度（用于断点续做判断）。

    返回:
      - completed_stages: 已完成的阶段列表
      - latest_stage: 最新完成的阶段
      - resume_stage: 下一步应执行的阶段（None 表示已完成）
      - is_active: 是否有未完成的任务
    """
    from app.services.working_memory import WorkingMemory

    wm = WorkingMemory(task_id)
    return {
        "task_id": task_id,
        "completed_stages": wm.get_completed_stages(),
        "latest_stage": wm.get_latest_stage(),
        "resume_stage": wm.get_resume_stage(),
        "is_active": wm.is_active(),
    }


@tasks_router.get("/tasks/{task_id}", response_model=TaskResponse)
async def get_task(task_id: str, user: GitHubUser | None = Depends(get_current_user)):
    task = get_session_manager().get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return TaskResponse(**task)


@tasks_router.get("/tasks/{task_id}/messages", response_model=list[MessageResponse])
async def get_task_messages(task_id: str, user: GitHubUser | None = Depends(get_current_user)):
    task = get_session_manager().get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return [MessageResponse(**m) for m in task.get("messages", [])]


@tasks_router.get("/tasks/{task_id}/files")
async def get_task_files(task_id: str, user: GitHubUser | None = Depends(get_current_user)):
    """任务文件区：上传的附件 + 生成的图表/结果文件。"""
    task = get_session_manager().get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {"task_id": task_id, "files": task.get("artifacts", [])}


@tasks_router.get("/tasks/{task_id}/events")
async def get_task_events(
    task_id: str,
    after: int = Query(0, ge=0, description="跳过前 N 条事件（增量回放）"),
    limit: int = Query(500, ge=1, le=2000, description="返回条数上限"),
    user: GitHubUser | None = Depends(get_current_user),
):
    """回放任务的持久化事件流（协议 v2.1）。

    事件按发生顺序写入 data/task_events/{task_id}.jsonl（node_start / node_end /
    plan / tool_call / tool_result / code_exec / task_end）。前端进会话时调用本端点
    恢复进度视图（dsh 式 session-projection 回放），配合 WS 实时事件使用。
    """
    from app.core.node_helpers import read_task_events

    events, total = read_task_events(task_id, after=after, limit=limit)
    return {
        "task_id": task_id,
        "events": events,
        "total": total,
        "after": after,
        "has_more": after + len(events) < total,
    }


@tasks_router.post("/tasks/{task_id}/cancel")
async def cancel_task(task_id: str, user: GitHubUser | None = Depends(get_current_user)):
    # cancel 含 sessions.json 全量落盘——绝不能在主循环线程里同步做
    # （曾与编排器线程争锁一起把事件循环堵死，所有请求 60s 无响应）
    success = await asyncio.to_thread(get_session_manager().cancel, task_id)
    if not success:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {"success": True, "message": "任务已取消"}


@tasks_router.get("/tasks", response_model=list[TaskResponse])
async def list_tasks():
    return [TaskResponse(**t) for t in get_session_manager().list_all()]


# ── Background orchestrator runner ────────────────────────────────


def _run_orchestrator_sync(task_id: str, problem: str, mode: str, user_id: str = "guest"):
    """在线程池中运行的同步入口（节点含阻塞调用，必须脱离事件循环）。"""
    try:
        # 子线程中 asyncio.run() 默认不创建 ThreadPoolExecutor，
        # 导致 langgraph 内部的 run_in_executor 调用失败。
        import concurrent.futures

        loop = asyncio.new_event_loop()
        loop.set_default_executor(concurrent.futures.ThreadPoolExecutor(max_workers=4))
        asyncio.set_event_loop(loop)
        loop.run_until_complete(_run_orchestrator(task_id, problem, mode, user_id))
    except Exception:
        import traceback

        with open("_orch_error.log", "a", encoding="utf-8") as f:
            f.write(f"\n=== {task_id} ===\n")
            traceback.print_exc(file=f)


async def _run_orchestrator(task_id: str, problem: str, mode: str, user_id: str = "guest"):
    """在后台运行 LangGraph 编排器。"""
    try:
        from app.core.nodes import TaskCancelledError, _is_cancelled
        from app.core.state import create_initial_state
        from app.core.workflow import get_orchestrator

        # 获取该用户的活跃 API Key
        active_key = get_active_api_key(user_id)

        orchestrator = get_orchestrator()
        state = create_initial_state(
            problem_raw=problem,
            mode=mode,
            session_id=task_id,
            api_key_config=active_key,
        )

        # 流式运行 — 每个节点完成后立即写入进度消息
        session_mgr = get_session_manager()
        publisher = None
        try:
            from app.services.redis_pubsub import get_publisher

            publisher = get_publisher()
        except Exception:
            publisher = None

        node_meta = {
            "classify_problem": ("问题分析", "识别问题类型"),
            "retrieve_knowledge": ("知识检索", "检索建模知识库"),
            "plan_execution": ("计划制定", "规划执行步骤"),
            "analysis_agent": ("问题分析", "深入剖析题意"),
            "modeling_agent": ("模型构建", "建立数学模型"),
            "solving_agent": ("求解计算", "编写并执行求解代码"),
            "verification_agent": ("验证分析", "检验模型鲁棒性"),
            "writing_agent": ("论文写作", "生成 LaTeX 论文"),
            "format_response": ("整合输出", "汇总最终结果"),
        }

        # node_name → 从 node_output 中取摘要的字段映射
        # 把"做了什么"摘要推给前端聊天面板，让用户看到每个 Agent 实际产出
        node_output_fields = {
            "analysis_agent": "analysis_output",
            "modeling_agent": "model_output",
            "solving_agent": "solving_output",
            "verification_agent": "verification_output",
            "writing_agent": "writing_output",
        }

        def _make_summary(node_name: str, node_output: dict) -> str:
            """从 node_output 中抽取该 Agent 的实际产出摘要（首 800 字）。"""
            field = node_output_fields.get(node_name)
            if not field:
                return ""
            text = node_output.get(field) or ""
            text = str(text).strip()
            if not text:
                return ""
            # 去掉前导 markdown 标题与多余空白
            text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
            text = re.sub(r"\s+", " ", text)
            return text[:800].strip()

        messages = []
        # 累积完整状态：updates 模式每个 chunk 只含当前节点返回的字段，
        # 逐块 update 合并才能得到完整 final state——否则只保留最后一块，
        # analysis/model/solving/verification 输出全部存成空串
        final_state = dict(state)

        async for chunk in orchestrator.astream(
            state, {"recursion_limit": 50}, stream_mode="updates"
        ):
            # 取消检查：任务被取消则提前退出编排器（节点入口另有检查）
            if _is_cancelled(task_id):
                raise TaskCancelledError(task_id)
            for node_name, node_output in chunk.items():
                stage, desc = node_meta.get(node_name, (node_name, f"执行: {node_name}"))
                summary = _make_summary(node_name, node_output) if node_output else ""

                # 构造更具体的聊天消息：标题 + 描述 + 摘要
                if summary:
                    content = f"[{stage}] {desc}\n\n{summary}{'…' if len(str(node_output.get(node_output_fields.get(node_name, ''), ''))) > 800 else ''}"
                else:
                    content = f"[{stage}] {desc}…"

                progress_msg = {
                    "id": str(uuid.uuid4())[:8],
                    "msg_type": "system",
                    "type": "info",
                    "content": content,
                    "agent_type": node_name.replace("_agent", "").replace("_", ""),
                    "created_at": None,
                }
                messages.append(progress_msg)
                session_mgr.update(task_id, messages=messages)
                # 实时推送到 WebSocket（fakeredis 或真实 Redis）
                # 前端拿 summary 渲染更具体的卡片，避免与顶部时间线信息重复
                if publisher:
                    publisher.node_end(
                        task_id,
                        node_name,
                        {
                            "stage": stage,
                            "title": stage,
                            "desc": desc,
                            "summary": summary,
                            "status": "completed",
                        },
                    )
                # messages 字段是各节点的增量列表，单独累积；
                # 其余字段直接 update（后写覆盖先写，各节点字段互不重叠）
                if node_output:
                    if node_output.get("messages"):
                        final_state.setdefault("messages", [])
                        final_state["messages"].extend(node_output["messages"])
                    final_state.update(
                        {k: v for k, v in node_output.items() if k != "messages"}
                    )

        result = final_state

        # 追加 agent 的详细信息到进度消息后面
        for msg in result.get("messages", []):
            messages.append(
                {
                    "id": str(uuid.uuid4())[:8],
                    "msg_type": msg.__class__.__name__.replace("Message", "").lower(),
                    "content": str(msg.content)[:500] if msg.content else None,
                    "created_at": None,
                }
            )

        session_mgr = get_session_manager()
        session_mgr.update(
            task_id,
            status="completed",
            final_response=result.get("final_response", ""),
            writing_output=result.get("writing_output", ""),
            analysis_output=result.get("analysis_output", ""),
            model_output=result.get("model_output", ""),
            solving_output=result.get("solving_output", ""),
            verification_output=result.get("verification_output", ""),
            messages=messages,
        )

        # 情景记忆：保存本次建模经验，供下次类似题召回
        try:
            from app.services.episodic_memory import EpisodicMemory

            em = EpisodicMemory()
            em.save(
                session_id=task_id,
                problem_raw=problem,
                problem_type=result.get("problem_type", ""),
                final_output=result.get("final_response", "") or result.get("writing_output", ""),
            )
        except Exception:
            pass  # 情景记忆不阻塞任务完成

        # 通知前端任务结束。
        # final_response 可能很长（论文含 LaTeX 全章），不在 WS payload 里塞全文：
        # WS 只推一个轻量级完成信号 + 前 800 字预览；前端再 GET /api/tasks/{id}
        # 拿到 writing_output / final_response 完整内容（已存到 session_mgr）。
        if publisher:
            publisher.task_end(
                task_id,
                "format_response",
                "completed",
                {
                    "final_response_preview": (result.get("final_response", "") or "")[:800],
                    "final_response_length": len(result.get("final_response", "") or ""),
                    "writing_output_length": len(result.get("writing_output", "") or ""),
                },
            )

    except TaskCancelledError:
        # 用户主动取消：标记任务为 cancelled（cancel 会置 status 并发出取消信号）
        logger.info("任务已被取消: %s", task_id)
        get_session_manager().cancel(task_id)
    except Exception as e:
        import traceback

        logger.error("Orchestrator failed for task %s:\n%s", task_id, traceback.format_exc())
        session_mgr = get_session_manager()
        session_mgr.update(
            task_id,
            status="error",
            final_response=f"错误: {str(e)}",
            messages=[
                {
                    "id": "error",
                    "msg_type": "system",
                    "content": f"主智能体运行失败: {str(e)}",
                }
            ],
        )
        try:
            from app.services.redis_pubsub import get_publisher

            get_publisher().task_end(task_id, "orchestrator", "error", {"message": str(e)})
        except Exception:
            pass
    finally:
        # 清理取消事件
        get_session_manager().cleanup_cancel_event(task_id)


# ── Document export ────────────────────────────────────────────────

import io

from fastapi import Query
from fastapi.responses import Response, StreamingResponse


@tasks_router.get("/tasks/{task_id}/export")
async def export_document(
    task_id: str,
    format: str = Query("md", description="导出格式: md | latex | docx | xlsx | csv"),
    user: GitHubUser | None = Depends(get_current_user),
):
    """导出方案模式生成的结果文档。

    - format=md:    Markdown 原文
    - format=latex: LaTeX .tex 文件
    - format=docx:  Word .docx 文件
    - format=xlsx:  Excel 结果表
    - format=csv:   CSV 数据文件
    """
    task = get_session_manager().get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    final_response = task.get("final_response", "")
    if not final_response:
        raise HTTPException(status_code=400, detail="任务尚未完成或未生成文档")

    # 尝试从 final_response 提取各部分
    writing = task.get("writing_output") or ""
    task.get("model_output") or ""

    if format == "latex":
        return _export_latex(final_response, writing)

    elif format == "docx":
        return _export_docx(final_response, title=task.get("problem", "数学建模方案")[:80])

    elif format == "xlsx":
        return _export_xlsx(task_id)

    elif format == "csv":
        return _export_csv(task_id)

    else:  # md
        return Response(
            content=final_response.encode("utf-8"),
            media_type="text/markdown; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename=modeling_solution_{task_id}.md",
            },
        )


def _export_latex(full_text: str, writing_output: str) -> Response:
    """从输出中提取/生成 LaTeX 并返回 .tex 下载。"""
    # 优先提取 writing 输出中的 LaTeX
    latex = writing_output or full_text

    # 尝试从 ```latex ... ``` 或 ```tex ... ``` 代码块中提取
    latex_match = re.search(r"```(?:latex|tex)?\s*\n(.*?)```", latex, re.DOTALL)
    if latex_match:
        latex = latex_match.group(1).strip()
    else:
        # 检查是否已经以 \documentclass 开头
        if not latex.strip().startswith("\\documentclass"):
            # 把全文当作 LaTeX-like 内容，包裹最小文档框架
            latex = (
                r"\documentclass[12pt,a4paper]{ctexart}"
                r"\usepackage{amsmath,amssymb,graphicx,booktabs,geometry}"
                r"\geometry{margin=2.5cm}"
                r"\begin{document}" + latex + r"\end{document}"
            )

    return Response(
        content=latex.encode("utf-8"),
        media_type="application/x-latex; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=modeling_paper.tex"},
    )


def _resolve_image_path(url: str) -> Path | None:
    """把 /api/task_files/... 或 /api/images/... URL 解析为本地文件路径。"""
    import tempfile as _tempfile

    parts = url.rstrip("/").split("/")
    settings = get_settings()
    try:
        if "task_files" in parts:
            # /api/task_files/{task_id}/{filename}
            task_id = parts[parts.index("task_files") + 1]
            filename = parts[-1]
            p = settings.project_root / "data" / "task_files" / task_id / filename
            if p.exists():
                return p
        elif "images" in parts:
            # /api/images/{run_id}/{filename} → temp 目录 + 持久副本兜底
            run_id = parts[parts.index("images") + 1]
            filename = parts[-1]
            for p in (
                Path(_tempfile.gettempdir()) / "mathmodel_outputs" / run_id / filename,
                settings.project_root / "data" / "chat_images" / run_id / filename,
            ):
                if p.exists():
                    return p
    except Exception:
        return None
    return None


def _export_docx(content: str, title: str = "数学建模方案") -> StreamingResponse:
    """将 Markdown 转换为 Word .docx 并返回（支持插入图表）。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # 标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _insert_image(url: str) -> None:
        """解析 markdown 图片语法并插入本地图表（居中，最大宽度 5.5 英寸）。"""
        try:
            path = _resolve_image_path(url)
            if not path:
                return
            doc.add_picture(str(path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass  # 图片插入失败不阻断导出

    # 简单 Markdown → docx 转换
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("```"):
            continue  # skip code fences

        # 图片：![说明](/api/task_files/... 或 /api/images/...) → 插入图表
        img_match = re.match(
            r"^!\[[^\]]*\]\((/api/(?:task_files|images)/[^)\s]+)\)$", stripped
        )
        if img_match:
            _insert_image(img_match.group(1))
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        elif stripped.startswith("|"):
            # 跳过表格（简化处理）
            continue
        elif stripped.startswith("$$") or stripped.startswith("$"):
            # 公式保留原文
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(11)
            run.italic = True
        else:
            p = doc.add_paragraph(stripped)

    # 输出到 buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=modeling_solution.docx"},
    )


def _export_xlsx(task_id: str) -> Response:
    """导出任务生成的 xlsx 结果文件。"""
    from pathlib import Path as _Path

    settings = get_settings()
    task_dir = _Path(settings.project_root) / "data" / "task_files" / task_id

    # 优先找 LLM 生成的 results.xlsx
    xlsx_files = sorted(task_dir.glob("*.xlsx"))
    if not xlsx_files:
        raise HTTPException(status_code=404, detail="未找到 xlsx 结果文件")

    xlsx_path = xlsx_files[0]
    return FileResponse(
        str(xlsx_path),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=f"results_{task_id}.xlsx",
    )


def _export_csv(task_id: str) -> Response:
    """导出任务生成的 CSV 数据文件。"""
    from pathlib import Path as _Path

    settings = get_settings()
    task_dir = _Path(settings.project_root) / "data" / "task_files" / task_id

    csv_files = sorted(task_dir.glob("*.csv"))
    if not csv_files:
        raise HTTPException(status_code=404, detail="未找到 CSV 数据文件")

    csv_path = csv_files[0]
    return FileResponse(
        str(csv_path),
        media_type="text/csv; charset=utf-8",
        filename=f"data_{task_id}.csv",
    )


# ── Package download ──────────────────────────────────────────────


@tasks_router.get("/tasks/{task_id}/package")
async def download_package(
    task_id: str,
    user: GitHubUser | None = Depends(get_current_user),
):
    """下载完整结果包（zip：论文 + 数据 + 图表 + 代码）。"""
    from pathlib import Path as _Path

    settings = get_settings()
    task_dir = _Path(settings.project_root) / "data" / "task_files" / task_id

    # 尝试找已有 zip
    zip_files = sorted(task_dir.glob("*.zip"))
    if zip_files:
        return FileResponse(
            str(zip_files[0]),
            media_type="application/zip",
            filename=f"{task_id}_results.zip",
        )

    # 动态生成 zip
    try:
        from app.services.result_packager import ResultPackager

        packager = ResultPackager(task_id, settings.project_root)
        zip_path = packager.build_zip_package()
        if zip_path.exists():
            return FileResponse(
                str(zip_path),
                media_type="application/zip",
                filename=f"{task_id}_results.zip",
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"打包失败: {e}")

    raise HTTPException(status_code=404, detail="无结果文件可打包")
