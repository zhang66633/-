"""导出接口 — Markdown → DOCX。

解析器对齐前端 marked 渲染的核心语法（标题/列表/表格/引用/代码块/链接/
图片/公式/加粗斜体删除线/分隔线），行内转换抽为纯函数 `_parse_inline`
（见 tests/test_export_md.py）。
"""

import io
import re

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

export_router = APIRouter()


class ExportDocxRequest(BaseModel):
    markdown: str
    title: str = "对话记录"


def _parse_inline(text: str) -> str:
    """把一行内的 Markdown 标记转成纯文本（与 marked 渲染后的可读内容对齐）。"""
    # 转义字符先占位保护（\* \_ \`），避免被下方格式正则误吞，最后还原
    text = text.replace(r"\*", "\x00A").replace(r"\_", "\x00B").replace(r"\`", "\x00C")
    # 行内代码
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # 图片 ![alt](url) → alt（图片: url）
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"\1（图片: \2）", text)
    # 链接 [text](url) → text（url）
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    # 加粗（** / __）
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    # 斜体（剩余单星）
    text = re.sub(r"\*([^*\n]+)\*", r"\1", text)
    # 删除线
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    # 公式：$$..$$ 与 $..$
    text = re.sub(r"\$\$(.+?)\$\$", r"[公式: \1]", text)
    text = re.sub(r"\$(.+?)\$", r"[\1]", text)
    # 还原转义字符
    text = text.replace("\x00A", "*").replace("\x00B", "_").replace("\x00C", "`")
    return text.strip()


_LIST_ITEM_RE = re.compile(r"^\s*(?:[-*+]|\d+[.)])\s+(.*)$")
_TASK_ITEM_RE = re.compile(r"^\s*[-*+]\s+\[([ xX])\]\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|.+\|\s*$")


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|[\s:|-]+\|\s*$", line))


def _add_table(doc, rows: list[str]) -> None:
    """把 GFM 表格行列表转成 docx 表格（跳过分隔行）。"""
    parsed = []
    for row in rows:
        if _is_table_separator(row):
            continue
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return
    n_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.style = "Table Grid"
    for ri, row in enumerate(parsed):
        for ci in range(n_cols):
            cell_text = _parse_inline(row[ci]) if ci < len(row) else ""
            table.cell(ri, ci).text = cell_text


def _add_list_item(doc, line: str) -> None:
    """列表/任务列表 → 缩进段落。"""
    from docx.shared import Pt

    m = _TASK_ITEM_RE.match(line)
    if m:
        prefix = "☑ " if m.group(1).lower() == "x" else "☐ "
        text = f"{prefix}{m.group(2)}"
    else:
        m = _LIST_ITEM_RE.match(line)
        if not m:
            return
        numbered = bool(re.match(r"^\s*\d+[.)]\s", line))
        prefix = "1. " if numbered else "• "
        text = f"{prefix}{m.group(1)}"
    p = doc.add_paragraph(_parse_inline(text))
    p.paragraph_format.left_indent = Pt(24)


def _markdown_to_docx(doc, markdown: str) -> None:
    """逐块解析 Markdown 并写入 docx。"""
    from docx.shared import Pt, RGBColor

    lines = markdown.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格（连续 | 行聚合成一个 docx table）
        if _TABLE_ROW_RE.match(line):
            table_rows = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                table_rows.append(lines[i])
                i += 1
            _add_table(doc, table_rows)
            continue

        # 标题
        if line.startswith("### "):
            doc.add_heading(_parse_inline(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(_parse_inline(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(_parse_inline(line[2:]), level=1)
        # 分隔线
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 40)
        # 引用
        elif line.startswith("> "):
            p = doc.add_paragraph(_parse_inline(line[2:]))
            p.paragraph_format.left_indent = Pt(24)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # 列表 / 任务列表
        elif _LIST_ITEM_RE.match(line):
            _add_list_item(doc, line)
        # 空行
        elif not line.strip():
            pass
        # 普通段落（行内标记 → 纯文本）
        else:
            doc.add_paragraph(_parse_inline(line))

        i += 1


@export_router.post("/export/docx")
async def export_docx(req: ExportDocxRequest):
    """将 Markdown 文本转为 .docx 文件下载。"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装")

    doc = Document()

    # 标题
    heading = doc.add_heading(req.title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _markdown_to_docx(doc, req.markdown)

    # 输出为流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{req.title}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
